import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import io

# --- FUNCIONES DE FORMATO ---
def agregar_numero_pagina(document):
    header = document.sections[0].header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.extend([fldChar1, instrText, fldChar2, fldChar3])

def procesar_docx(file_upload, titulo, estudiantes, asignatura, docente, fecha):
    doc_original = Document(file_upload)
    nuevo_doc = Document()
    
    for section in nuevo_doc.sections:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    
    agregar_numero_pagina(nuevo_doc)

    # Portada (Distribución compacta para celular/PC)
    p_titulo = nuevo_doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_titulo.add_run(titulo.upper())
    run_t.bold = True
    run_t.font.name = 'Times New Roman'
    run_t.font.size = Pt(12)
    
    for _ in range(4): nuevo_doc.add_paragraph()
    nuevo_doc.add_paragraph("Presentado por:").alignment = WD_ALIGN_PARAGRAPH.CENTER
    estudiantes.sort()
    for est in estudiantes:
        nuevo_doc.add_paragraph(est.strip()).alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(4): nuevo_doc.add_paragraph()
    nuevo_doc.add_paragraph(asignatura).alignment = WD_ALIGN_PARAGRAPH.CENTER
    nuevo_doc.add_paragraph(docente).alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(3): nuevo_doc.add_paragraph()
    nuevo_doc.add_paragraph("Universidad de Cundinamarca").alignment = WD_ALIGN_PARAGRAPH.CENTER
    nuevo_doc.add_paragraph("Administración de empresas").alignment = WD_ALIGN_PARAGRAPH.CENTER
    nuevo_doc.add_paragraph("Sede Facatativá").alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(3): nuevo_doc.add_paragraph()
    nuevo_doc.add_paragraph(fecha).alignment = WD_ALIGN_PARAGRAPH.CENTER

    nuevo_doc.add_page_break()

    # Cuerpo
    for para in doc_original.paragraphs:
        if para.text.strip():
            nuevo_p = nuevo_doc.add_paragraph(para.text)
            nuevo_p.paragraph_format.first_line_indent = Inches(0.5)
            nuevo_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            for run in nuevo_p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

    target = io.BytesIO()
    nuevo_doc.save(target)
    return target.getvalue()

# --- INTERFAZ WEB ---
st.set_page_config(page_title="Generador APA UDEC", page_icon="📝")
st.title("📝 Formateador APA 7 - UDEC")
st.info("Sube tu Word y obtén la versión con normas APA 7 instantáneamente.")

archivo = st.file_uploader("Sube tu archivo Word (.docx)", type=["docx"])
titulo = st.text_input("Título del trabajo")
est_raw = st.text_input("Estudiantes (separados por coma)")
asig = st.text_input("Asignatura")
prof = st.text_input("Docente")

if st.button("Generar Documento APA"):
    if archivo and titulo and est_raw:
        meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
        hoy = datetime.datetime.now()
        fecha_es = f"{hoy.day} de {meses[hoy.month-1]} del {hoy.year}"
        
        resultado = procesar_docx(archivo, titulo, est_raw.split(","), asig, prof, fecha_es)
        
        st.download_button(
            label="📥 Descargar Word Formateado",
            data=resultado,
            file_name=f"{titulo}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.success("¡Documento generado con éxito!")
    else:
        st.warning("Por favor completa todos los campos.")